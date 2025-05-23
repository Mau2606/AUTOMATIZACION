import shutil
from docx import Document

class WordHandler:
    def __init__(self, plantilla_path):
        self.plantilla_path = plantilla_path

    def copiar_y_renombrar(self, nuevo_path):
        shutil.copy(self.plantilla_path, nuevo_path)

    def reemplazar_en_runs(self, parrafo, marcador, valor):
        """
        Reemplaza un marcador en cada run del párrafo.
        """
        for run in parrafo.runs:
            if marcador in run.text:
                run.text = run.text.replace(marcador, valor)

    def reemplazar_con_formato(self, parrafo, marcador, valor):
        """
        Reemplaza el marcador en un párrafo y aplica formato.
        """
        for run in parrafo.runs:
            if marcador in run.text:
                run.text = run.text.replace(marcador, valor)
                # Aplicar formato Arial 11, negrita, subrayado
                run.font.name = 'Arial'
                #run.font.size = 11
                run.bold = True
                run.underline = True


    def rellenar_datos_completos(self, ruta_word, titulo, fecha_en_palabras, ruc, rit,  juez, ct, audio, sala, inicio):
        doc = Document(ruta_word)
        # Reemplazar el título con formato
        for parrafo in doc.paragraphs:
            if '{TITULO}' in parrafo.text:
                self.reemplazar_con_formato(parrafo, '{TITULO}', titulo)
            if '{RIT}' in parrafo.text:
                self.reemplazar_en_runs(parrafo, '{RIT}', rit)
            if '{RUC}' in parrafo.text:
                self.reemplazar_en_runs(parrafo, '{RUC}', ruc)
            if '{JUEZ}' in parrafo.text:
                self.reemplazar_en_runs(parrafo, '{JUEZ}', juez.title())


        # Reemplazar los marcadores en las tablas del documento
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        self.reemplazar_en_runs(parrafo, '{FECHA}', fecha_en_palabras)
                        self.reemplazar_en_runs(parrafo, '{RUC}', ruc)
                        self.reemplazar_en_runs(parrafo, '{RIT}', rit)
                        self.reemplazar_en_runs(parrafo, '{JUEZ}', juez)
                        self.reemplazar_en_runs(parrafo, '{CT}', ct)
                        self.reemplazar_en_runs(parrafo, '{AUDIO}', audio)
                        self.reemplazar_en_runs(parrafo, '{SALA}', sala)
                        self.reemplazar_en_runs(parrafo, '{INICIO}', inicio)


        # Guardar el archivo con los datos reemplazados
        doc.save(ruta_word)